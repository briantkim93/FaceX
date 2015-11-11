__author__ = 'Kim'

import os
import hashlib
import time
import datetime
import cv2
import facex
import xlsxwriter
import docx


temp_pic_dir = "temp.data/"
face_dir = "face.data/"
bin_dir = "bin/"
face_bin_arch = "%s/eigen.bin" % bin_dir
system_user_dir = "C:/Users/%s/Pictures" % os.environ.get("USERNAME")
system_user_doc_dir = "C:/Users/%s/Documents/FaceX" % os.environ.get("USERNAME")


def ensure_data():
    if not os.path.isdir(system_user_doc_dir):
        os.mkdir(system_user_doc_dir)

    if not os.path.isdir(temp_pic_dir):
        os.mkdir(temp_pic_dir)

    if not os.path.isdir(face_dir):
        os.mkdir(face_dir)

    if not os.path.isdir(bin_dir):
        os.mkdir(bin_dir)

    f = open(face_bin_arch, "ab")
    f.close()


def atexit():
    print "exiting now"
    # delete temporary data
    entries = os.listdir(temp_pic_dir)
    for entry in entries:
        fullpath = "%s/%s" % (temp_pic_dir, entry)
        os.unlink(fullpath)

    os.rmdir(temp_pic_dir)


def get_random_name():
    epoch = time.time()
    encry = hashlib.md5(str(epoch))
    name = "%s.png" % encry.hexdigest()
    return name


def is_emptystr(inp):
    if len(str(inp)) == 0:
        return True
    return False


def is_emptystr_group(inp):
    for data in inp:
        if is_emptystr(data):
            return True

    return False


def stamp_to_date(stamp):
    date = datetime.datetime.fromtimestamp(int(stamp)).strftime("%d-%m-%Y %H:%M:%S")
    return date

def date_to_stamp(date):
    stamp = time.mktime(datetime.datetime.strptime(date, "%d/%m/%Y").timetuple())
    return stamp

def getaccessdocname(doc_type):
    ltime = time.gmtime()
    doc_name = "%d%d%d%d%d%d" % (ltime.tm_year, ltime.tm_mon, ltime.tm_mday, ltime.tm_hour, ltime.tm_min, ltime.tm_sec)

    doc_name = "%s.%s" % (doc_name, doc_type)

    return doc_name

def temp_to_permanent_face(temps):
    """

    :rtype : list
    """
    perm = []

    for temp in temps:
        src = "%s%s" % (temp_pic_dir, temp)
        dest = "%s%s" % (face_dir, temp)

        mat = cv2.imread(src)
        cv2.imwrite(dest, mat)

        perm.append(dest)

    return perm

def publish_xlsx_report(d_from, d_to):
    s_from = date_to_stamp(d_from)
    s_to = date_to_stamp(d_to)

    datalist = facex.datacenter.get_access_reports(s_from, s_to)
    fname = getaccessdocname(doc_type="xlsx")
    fpath = system_user_doc_dir + "/" + fname
    workbook = xlsxwriter.Workbook(fpath)
    worksheet = workbook.add_worksheet("Reports")
    c_row = 1

    worksheet.write(0, 0, "DATE")
    worksheet.write(0, 1, "ID NUMBER")
    worksheet.write(0, 2, "ACCESS STATUS")

    for data in datalist:
        worksheet.write(c_row, 0, data["Date"])
        worksheet.write(c_row, 1, data["id_no"])
        worksheet.write(c_row, 2, data["Access"])

        c_row += 1

    workbook.close()

    return fpath

def publish_docx_report(d_from, d_to):

    s_from = date_to_stamp(d_from)
    s_to = date_to_stamp(d_to)

    datalist = facex.datacenter.get_access_reports(s_from, s_to)

    fname = getaccessdocname(doc_type="docx")
    fpath = system_user_doc_dir + "/" + fname

    book = docx.Document()
    book.add_heading("FACEX System Access Report", 0)
    from_date = stamp_to_date(s_from)
    to_date = stamp_to_date(s_to)

    book.add_heading("FACEX System Access Report", level=1)
    book.add_heading("From " + from_date + " To " + to_date, level=3)

    table = book.add_table(rows=len(datalist) + 1, cols=3)
    header_cells = table.rows[0].cells
    header_cells[0].text = "DATE"
    header_cells[1].text = "ID NUMBER"
    header_cells[2].text = "ACCESS"

    for data in datalist:
        row_cells = table.add_row().cells
        row_cells[0].text = data["Date"]
        row_cells[1].text = data["id_no"]
        row_cells[2].text = data["Access"]

    book.add_page_break()
    book.save(fpath)

    return fpath

def publish_xlsx_userlist():
    datalist = facex.datacenter.get_allmembers()
    fname = getaccessdocname(doc_type="xlsx")
    fpath = system_user_doc_dir + "/" + fname
    workbook = xlsxwriter.Workbook(fpath)
    worksheet = workbook.add_worksheet("Members")
    c_row = 1

    worksheet.write(0, 0, "NAMES")
    worksheet.write(0, 1, "ID NUMBER")
    worksheet.write(0, 2, "DEPARTMENT")
    worksheet.write(0, 3, "POSITION")

    for data in datalist:
        worksheet.write(c_row, 0, data.names)
        worksheet.write(c_row, 1, data.id_no)
        worksheet.write(c_row, 2, data.department)
        worksheet.write(c_row, 3, data.position)

        c_row += 1

    workbook.close()

    return fpath

def publish_docx_userlist():
    datalist = facex.datacenter.get_allmembers()

    fname = getaccessdocname(doc_type="docx")
    fpath = system_user_doc_dir + "/" + fname

    book = docx.Document()
    book.add_heading("FACEX System Members", 0)

    book.add_heading("FACEX System Members list", level=1)

    table = book.add_table(rows=len(datalist) + 1, cols=4)
    header_cells = table.rows[0].cells
    header_cells[0].text = "NAMES"
    header_cells[1].text = "ID NUMBER"
    header_cells[2].text = "DEPARTMENT"
    header_cells[3].text = "POSITION"

    for data in datalist:
        row_cells = table.add_row().cells
        row_cells[0].text = data.names
        row_cells[1].text = str(data.id_no)
        row_cells[2].text = data.department
        row_cells[3].text = data.position

    book.add_page_break()
    book.save(fpath)

    return fpath
